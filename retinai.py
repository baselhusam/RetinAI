import streamlit as st
import numpy as np
from PIL import Image
import os
import base64
from docx2pdf import convert
import pypdfium2 as pdfium
import time

from keras import layers
from keras.applications import DenseNet121
from keras.models import Sequential
from keras.optimizers import Adam

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches
from datetime import date

st.set_page_config(
    page_title="RetinAI",
    page_icon="assets/icon.png",
    layout="centered", )

def new_line(n=1):
    for _ in range(n):
        st.write("")

@st.cache_data
def build_model():

    densenet = DenseNet121(weights='imagenet',
                           include_top=False,
                           input_shape=(224,224,3))
            
    model = Sequential()
    model.add(densenet)
    model.add(layers.GlobalAveragePooling2D())
    model.add(layers.Dropout(0.5))
    model.add(layers.Dense(5, activation='softmax'))
    
    model.compile(
        loss='categorical_crossentropy',
        optimizer=Adam(learning_rate=0.00005),
        metrics=['accuracy']
    )
    
    return model

model = build_model()

def predict(image, model= model):
    model.load_weights("Y_new_1.h5")

    image = np.asarray(Image.open(image).resize((224,224)))
    image = image.reshape(1,224,224,3)
    image = image/255.0

    pred = model.predict(image)
    pred = pred.argmax(axis=1)
    return pred

def processing_prediction(pred):

    pred_dict = {0: "No DR", 1: "Mild", 2: "Moderate", 3: "Severe", 4: "Proliferative DR"}

    return pred_dict[pred[0]]

def save_uploadedfile(uploadedfile, output_name):
    with open(output_name, "wb") as f:
        f.write(uploadedfile.getbuffer())


# -------- Create the Report --------

def create_replacements_dict(input_data):

    """
    This function creates a dictionary of the input data to be used for replacing the text in the docx template file.
    
    Parameters
    ----------
    input_data : dict
        A list containing the input data.

    Returns
    -------
    replacements : dict
        A dictionary containing the input data in the format required for replacing the text in the docx template file.
    """

    replacements = {
        "<PATIENT_ID>": "0000",
        "<PATIENT_NAME>": "Rula",
        "<PATIENT_DOB>": "1/1/1999",
        "<PATIENT_GENDER>": "Female",
        "<LOCATION>": "Amman Hospital",
        "<PROVIDER>": "Dr. Tareq",
        "<CONTROL_ID>": "1111",
        "<DATE>": str(date.today()),
        "<LEFT_IMAGE_1>": input_data[0],
        "<RIGHT_IMAGE_1>":input_data[1],
        "<LEFT_CLASS>":input_data[2],
        "<RIGHT_CLASS>":input_data[3]
    }

    if input_data[2] == "Proliferative DR" or input_data[3] =="Proliferative DR":
       replacements["<RECOMMENDATION>"]= "This is the advanced stage where fragile new blood vessels grow but can leak blood, risking vision. Immediate treatment is essential as PDR can cause severe vision complications, including blindness. Regular monitoring and potential surgical interventions are crucial at this stage."
    elif input_data[2] == "Severe" or input_data[3] =="Severe":
      replacements["<RECOMMENDATION>"]= "At this stage, more blood vessels are becoming blocked, leading to areas of the retina being deprived of blood. As a result, growth factors are secreted, signaling the growth of new blood vessels. Due to the severity, immediate consultation with a retinal specialist is advised, along with strict blood sugar control."
    elif input_data[2] == "Moderate" or input_data[3] =="Moderate":
      replacements["<RECOMMENDATION>"]= "At this stage, the nourishing blood vessels of the retina may swell, distort, or lose their ability to transport blood. It's essential to regularly monitor the eyes and consult with healthcare providers to optimize diabetes management."
    elif input_data[2] == "Mild" or input_data[3] =="Mild":
      replacements["<RECOMMENDATION>"]= "This is the initial stage of diabetic retinopathy, marked by tiny areas of swelling in the retina's blood vessels, known as microaneurysms. To manage this, it's advised to schedule more frequent eye exams and ensure control over blood sugar levels, blood pressure, and cholesterol."
    else:
        replacements["<RECOMMENDATION>"]= "There are no signs of retinal damage associated with diabetes. While your eyes are currently healthy, it's crucial to continue regular eye check-ups as recommended and maintain blood sugar levels to e sure continued eye health."

    return replacements


def replace_text_in_table(table, replacements):

    """
    This function replaces the text in the docx template file with the input data.

    Parameters
    ----------
    table : docx.table
        The table in the docx template file.
    replacements : dict
        A dictionary containing the input data in the format required for replacing the text in the docx template file.

    Returns
    -------
    None.

    Note: This function is specific to the docx template file. It needs to be modified if the template file is changed.
    """


    #Replace Text
    for row in table.rows:
        for cell in row.cells:
            for key, value in replacements.items():
                if key in cell.text:
                    cell.text = cell.text.replace(key, value)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(55, 53, 69)
    #Replace Images
    table.rows[9].cells[0]._element.clear_content()
    img = table.rows[9].cells[0].add_paragraph().add_run().add_picture(replacements["<LEFT_IMAGE_1>"], width=Inches(1.64))

    table.rows[9].cells[4]._element.clear_content()
    img = table.rows[9].cells[4].add_paragraph().add_run().add_picture(replacements["<RIGHT_IMAGE_1>"], width=Inches(1.64))


def edit_docx_template(template_file, output_filename, input_data, output_file):

    """
    Edit the docx template file with the input data and save it as a new file.

    Parameters
    ----------
    template_file : str
        The path to the docx template file.
    output_filename : str
        The name of the output file.
    input_data : dict
        A dictionary containing the input data.
    output_file : str
        The path to the output file.

    Returns
    -------
    None.

    """

    replacements = create_replacements_dict(input_data)

    table = template_file.tables
    replace_text_in_table(table[0], replacements)

    template_file.save(output_file)

left_pred = None
right_pred = None
input_data = []

col1, col2, col3 = st.columns([0.5,2,0.5])
col2.image("assets/logo.png")
new_line(2)

st.write("""
**RetinAI** is a web application that allows you to upload your retinal images and get a prediction of the disease you may have.
We use a deep learning technology to predict the level of the disease that you have with 85% accuracy.
""")

st.divider()

st.write(""" 
For the Diabetes Retinopathy, we can predict the level of the following classes:

**1. No DR:** No diabetic retinopathy <br>
**2. Mild:** Mild nonproliferative retinopathy <br>
**3. Moderate:** Moderate nonproliferative retinopathy <br>
**4. Severe:** Severe nonproliferative retinopathy <br>
**5. Proliferative DR:** Proliferative retinopathy <br>

""", unsafe_allow_html=True)
new_line(3)


st.markdown("<h4 align='center'>Upload your retinal images</h4>", unsafe_allow_html=True)
new_line(2)

col1, col2 = st.columns(2, gap='medium')

with col1:
    st.markdown("<h6 align='center'>Upload your Left retinal image</h5>", unsafe_allow_html=True)
    left_image = st.file_uploader("Upload your Left retinal image", type=['png', 'jpg', 'jpeg'])

    if left_image is not None:
        st.image(left_image, use_column_width=True)
        save_uploadedfile(left_image, "left_image.png")


with col2:
    st.markdown("<h6 align='center'>Upload your Right retinal image</h5>", unsafe_allow_html=True)
    right_image = st.file_uploader("Upload your Right retinal image", type=['png', 'jpg', 'jpeg'])

    if right_image is not None:
        st.image(right_image, use_column_width=True)
        save_uploadedfile(right_image, "right_image.png")

new_line(2)
cola, colb, colc = st.columns([1,.5,1])
if colb.button("🔮 Predict", use_container_width=True):
    if left_image is not None and right_image is not None:
        
        left_pred = predict(left_image)
        right_pred = predict(right_image)

        col1.info("Left eye has: **" + processing_prediction(left_pred)+ "**" )
        col2.info("Right eye has: **" + processing_prediction(right_pred) + "**")

        input_data = ["left_image.png", "right_image.png", processing_prediction(left_pred), processing_prediction(right_pred)]

    else:
        st.error("Please upload both images to predict the disease.")

# Check if the Input is ready to make and show the report to bee seen and downloaded
if input_data:

    # Create Docx Report
    template_file = Document("RetinAI_Report_Template.docx")
    output_file = r"RetinAI_Report.docx"

    edit_docx_template(template_file, output_file, input_data, output_file)

    import pythoncom
    from docx2pdf import convert
    pythoncom.CoInitialize()

    convert(output_file)


    st.markdown("<h4 align='center'>RetinAI Report</h4>", unsafe_allow_html=True)

    # Convert PDF to Image    
    filepath = "RetinAI_Report.pdf"
    pdf = pdfium.PdfDocument(filepath)

    page = pdf[0]
    pil_image = page.render(scale=4).to_pil()
    pil_image.save("RetinAI_Report_Image.jpg")

    st.image("RetinAI_Report_Image.jpg", use_column_width=True)

    # Download the report
    col1, col2, col3 = st.columns([1,1,1])
    if col2.download_button("📥 Download Report", 
                            data=open(filepath, 'rb').read(), 
                            file_name="RetinAI_Report.pdf", 
                            mime="application/pdf",):
        
        pass

    
